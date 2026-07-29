"""DOCX 文本提取器（从 api.py 提取的逻辑）"""
import io
from docx import Document


def extract_docx_text(raw: bytes) -> str:
    """从 DOCX 文件提取文本（包括段落和表格）"""
    doc = Document(io.BytesIO(raw))
    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts)
